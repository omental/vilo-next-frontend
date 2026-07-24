import hashlib
from datetime import date, datetime, timedelta, timezone
from io import BytesIO
from pathlib import Path
from typing import Any

import httpx
from fastapi import APIRouter, BackgroundTasks, Depends, File, Form, HTTPException, Request, UploadFile, status
from fastapi.responses import FileResponse
from jose import JWTError, jwt
from sqlalchemy import String, and_, cast, func, or_, select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.api.deps import role_guard
from app.core.config import settings
from app.db.session import get_db
from app.models.case import Case
from app.models.client import Client
from app.models.document import Document
from app.models.document_version import DocumentVersion
from app.models.user import User
from app.schemas.document import (
    DocumentEditableContentResponse,
    DocumentEditableContentUpdate,
    DocumentListResponse,
    DocumentResponse,
    DocumentUpdate,
    DocumentVersionResponse,
    OnlyOfficeSessionResponse,
)
from app.services.audit import log_audit_event
from app.services.document_storage import MAX_UPLOAD_BYTES, persist_file, resolve_stored_file, safe_original_name
from app.services.email import build_document_shared_email
from app.services.jobs import enqueue_email
from app.services.notifications import create_notification
from app.services.timeline import create_case_timeline_event
from app.services.access import accessible_case_condition

router = APIRouter(prefix="/documents", tags=["documents"])
ALLOWED_STAFF = ["partner", "admin", "lawyer", "paralegal"]
VALID_VISIBILITY = {"internal", "client_visible"}
STORAGE_ROOT = Path("backend/storage/documents")
DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
DOCX_WARNING = "Editing DOCX content creates a new version. Original uploaded file remains in version history."
ONLYOFFICE_CALLBACK_SUCCESS = {"error": 0}
ONLYOFFICE_SAVE_STATUSES = {2, 6}


def to_response(document: Document) -> DocumentResponse:
    loaded = getattr(document, "__dict__", {})
    return DocumentResponse(
        id=document.id,
        organization_id=document.organization_id,
        case_id=document.case_id,
        client_id=document.client_id,
        uploaded_by=document.uploaded_by,
        title=document.title,
        description=document.description,
        file_name=document.file_name,
        file_type=document.file_type,
        file_size=document.file_size,
        category=document.category,
        visibility=document.visibility,
        version=document.version,
        version_source=getattr(document, "version_source", None),
        version_note=getattr(document, "version_note", None),
        created_at=document.created_at,
        updated_at=document.updated_at,
        case_title=getattr(loaded.get("case"), "title", None),
        client_name=getattr(loaded.get("client"), "name", None),
        uploader_name=getattr(loaded.get("uploader"), "name", None),
    )


def to_version_response(version: DocumentVersion) -> DocumentVersionResponse:
    return DocumentVersionResponse(
        id=version.id,
        document_id=version.document_id,
        organization_id=version.organization_id,
        file_name=version.file_name,
        file_type=version.file_type,
        file_size=version.file_size,
        version_number=version.version_number,
        uploaded_by=version.uploaded_by,
        source=getattr(version, "source", None),
        notes=version.notes,
        version_note=version.notes,
        created_at=version.created_at,
    )


async def get_org_document(db: AsyncSession, document_id: int, organization_or_user: int | User) -> Document | None:
    organization_id = organization_or_user.organization_id if isinstance(organization_or_user, User) else organization_or_user
    query = select(Document).outerjoin(Case, Case.id == Document.case_id).where(
        Document.id == document_id,
        Document.organization_id == organization_id,
    ).options(selectinload(Document.case), selectinload(Document.client), selectinload(Document.uploader))
    if isinstance(organization_or_user, User):
        query = query.where(or_(Document.case_id.is_(None), accessible_case_condition(organization_or_user)))
    return await db.scalar(query)


def is_docx_document(document: Document) -> bool:
    file_name = (document.file_name or "").lower()
    file_type = (document.file_type or "").lower()
    return file_name.endswith(".docx") or file_type == DOCX_MIME_TYPE


def extract_docx_text(file_path: str) -> str:
    from docx import Document as DocxDocument

    path = Path(file_path)
    if not path.exists():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Stored file not found")

    docx_document = DocxDocument(str(path))
    blocks: list[str] = []

    for paragraph in docx_document.paragraphs:
        text = paragraph.text.rstrip()
        blocks.append(text)

    for table_index, table in enumerate(docx_document.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            if blocks and any(blocks):
                blocks.append("")
            blocks.append(f"[Table {table_index}]")
            blocks.extend(rows)

    text = "\n".join(blocks).strip()
    return text


def render_docx_bytes(content: str) -> bytes:
    from docx import Document as DocxDocument

    docx_document = DocxDocument()
    normalized = (content or "").replace("\r\n", "\n")
    lines = normalized.split("\n")

    if not any(line.strip() for line in lines):
        docx_document.add_paragraph("")
    else:
        for line in lines:
            docx_document.add_paragraph(line)

    buffer = BytesIO()
    docx_document.save(buffer)
    return buffer.getvalue()


def build_docx_version_name(document: Document) -> str:
    original_name = (document.file_name or "").strip()
    if original_name.lower().endswith(".docx"):
        return original_name
    title = (document.title or "").strip()
    stem = title or original_name or f"document-{document.id}"
    safe_stem = "".join(char if char.isalnum() or char in {"-", "_", " "} else "-" for char in stem).strip(" -_")
    if not safe_stem:
        safe_stem = f"document-{document.id}"
    safe_stem = safe_stem.replace(" ", "-")
    return f"{safe_stem}.docx"


def archive_current_document_version(doc: Document, actor_user_id: int, now: datetime) -> DocumentVersion:
    return DocumentVersion(
        document_id=doc.id,
        organization_id=doc.organization_id,
        file_name=doc.file_name,
        file_path=doc.file_path,
        file_type=doc.file_type,
        file_size=doc.file_size,
        version_number=doc.version,
        uploaded_by=actor_user_id,
        source=getattr(doc, "version_source", "upload"),
        notes=getattr(doc, "version_note", None),
        created_at=now,
    )


def current_utc() -> datetime:
    return datetime.now(timezone.utc)


def normalize_base_url(value: str) -> str:
    return value.rstrip("/")


def get_public_backend_base_url(request: Request) -> str:
    configured = (settings.public_backend_url or "").strip()
    if configured:
        return normalize_base_url(configured)
    return normalize_base_url(str(request.base_url).rstrip("/"))


def build_onlyoffice_document_key(doc: Document) -> str:
    digest = hashlib.sha256(f"{doc.organization_id}:{doc.id}:{doc.version}:{doc.updated_at}".encode("utf-8")).hexdigest()
    return digest[:48]


def build_internal_document_token(*, document_id: int, organization_id: int, version: int, purpose: str) -> str:
    expires_at = current_utc() + timedelta(minutes=settings.onlyoffice_file_token_expires_minutes)
    payload = {
        "sub": f"document:{document_id}",
        "document_id": document_id,
        "organization_id": organization_id,
        "version": version,
        "purpose": purpose,
        "exp": expires_at,
    }
    return jwt.encode(payload, settings.jwt_secret_key, algorithm=settings.jwt_algorithm)


def decode_internal_document_token(token: str, expected_purpose: str) -> dict[str, Any]:
    try:
        payload = jwt.decode(token, settings.jwt_secret_key, algorithms=[settings.jwt_algorithm])
    except JWTError as exc:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Invalid or expired ONLYOFFICE token") from exc
    if payload.get("purpose") != expected_purpose:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Invalid ONLYOFFICE token purpose")
    return payload


def sign_onlyoffice_config(config: dict[str, Any]) -> dict[str, Any]:
    secret = (settings.onlyoffice_jwt_secret or "").strip()
    if not secret:
        return config
    signed = dict(config)
    signed["token"] = jwt.encode(config, secret, algorithm="HS256")
    return signed


def decode_onlyoffice_callback_token(token: str) -> dict[str, Any]:
    secret = (settings.onlyoffice_jwt_secret or "").strip()
    if not secret:
        return {}
    try:
        return jwt.decode(token, secret, algorithms=["HS256"])
    except JWTError as exc:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Invalid ONLYOFFICE callback JWT") from exc


def get_onlyoffice_callback_payload(request: Request, body: dict[str, Any]) -> dict[str, Any]:
    secret = (settings.onlyoffice_jwt_secret or "").strip()
    if not secret:
        return body

    auth_header = request.headers.get("authorization", "")
    bearer_token = auth_header.removeprefix("Bearer ").strip() if auth_header.startswith("Bearer ") else ""
    body_token = body.get("token") if isinstance(body.get("token"), str) else ""
    token = bearer_token or body_token
    if not token:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="ONLYOFFICE callback JWT is required")

    decoded = decode_onlyoffice_callback_token(token)
    if bearer_token:
        return body
    return decoded


def build_onlyoffice_callback_note(payload: dict[str, Any]) -> str:
    users = payload.get("users")
    if isinstance(users, list) and users:
        return f"Edited in ONLYOFFICE by {users[0]}"
    return "Edited in ONLYOFFICE"


def get_onlyoffice_actor_user_id(payload: dict[str, Any], fallback_user_id: int) -> int:
    users = payload.get("users")
    if isinstance(users, list) and users:
        try:
            return int(users[0])
        except (TypeError, ValueError):
            return fallback_user_id
    return fallback_user_id


async def validate_case(db: AsyncSession, current_user: User, case_id: int | None):
    if case_id is None:
        return None
    case = await db.scalar(select(Case).where(Case.id == case_id, Case.organization_id == current_user.organization_id, accessible_case_condition(current_user)))
    if not case:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Case must belong to your organization")
    return case


async def validate_client(db: AsyncSession, organization_id: int, client_id: int | None):
    if client_id is None:
        return None
    client = await db.scalar(select(Client).where(Client.id == client_id, Client.organization_id == organization_id))
    if not client:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Client must belong to your organization")
    return client


@router.post("/upload", response_model=DocumentResponse)
async def upload_document(
    title: str = Form(...),
    description: str | None = Form(default=None),
    category: str | None = Form(default=None),
    visibility: str = Form(default="internal"),
    case_id: int | None = Form(default=None),
    client_id: int | None = Form(default=None),
    file: UploadFile = File(...),
    request: Request = None,
    background_tasks: BackgroundTasks = None,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(role_guard(ALLOWED_STAFF)),
):
    if visibility not in VALID_VISIBILITY:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid visibility")
    case = await validate_case(db, current_user, case_id)
    client = await validate_client(db, current_user.organization_id, client_id)
    if case is not None and client is not None and case.client_id != client.id:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Case does not belong to client")
    resolved_client_id = client.id if client is not None else (case.client_id if case is not None else None)

    original_name = safe_original_name(file.filename or "")
    data = await file.read()
    if not data:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Empty file")
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="File exceeds upload size limit")

    file_path, _stored_name = persist_file(STORAGE_ROOT, current_user.organization_id, original_name, data)

    now = datetime.now(timezone.utc)
    document = Document(
        organization_id=current_user.organization_id,
        case_id=case_id,
        client_id=resolved_client_id,
        uploaded_by=current_user.id,
        title=title,
        description=description,
        file_name=original_name,
        file_path=str(file_path),
        file_type=file.content_type,
        file_size=len(data),
        category=category,
        visibility=visibility,
        version=1,
        version_source="upload",
        version_note=None,
        created_at=now,
        updated_at=now,
    )
    db.add(document)
    await db.flush()
    await log_audit_event(
        db,
        organization_id=current_user.organization_id,
        user_id=current_user.id,
        action="document_uploaded",
        entity_type="document",
        entity_id=str(document.id),
        description=f"Document uploaded: {document.title}",
        metadata_json={"case_id": document.case_id, "visibility": document.visibility},
        ip_address=request.client.host if request and request.client else None,
        user_agent=request.headers.get("user-agent") if request else None,
    )

    if case is not None:
        await create_case_timeline_event(
            db,
            organization_id=current_user.organization_id,
            case_id=case.id,
            actor_id=current_user.id,
            event_type="document_uploaded",
            title=f"Document uploaded: {title}",
            metadata_json={"document_id": document.id, "file_name": original_name},
        )
    if visibility == "client_visible" and case is not None:
        client = await db.scalar(select(Client).where(Client.id == case.client_id, Client.organization_id == current_user.organization_id))
        if client and client.user_id:
            await create_notification(
                db,
                organization_id=current_user.organization_id,
                user_id=client.user_id,
                type="document_shared",
                title=f"New shared document: {title}",
                body="A document has been shared with you in the client portal.",
                metadata_json={"document_id": document.id, "case_id": case.id},
            )
            if background_tasks and client.email:
                subject, html_body, text_body = build_document_shared_email(
                    client_name=client.name or "Client",
                    document_title=document.title,
                    document_id=document.id,
                )
                enqueue_email(background_tasks, to_email=client.email, subject=subject, html_body=html_body, text_body=text_body)

    await db.commit()
    await db.refresh(document)
    return to_response(document)


@router.get("", response_model=list[DocumentResponse])
async def list_documents(
    case_id: int | None = None,
    client_id: int | None = None,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(role_guard(ALLOWED_STAFF)),
):
    query = select(Document).outerjoin(Case, Case.id == Document.case_id).where(
        Document.organization_id == current_user.organization_id,
        or_(Document.case_id.is_(None), accessible_case_condition(current_user)),
    )
    if case_id is not None:
        query = query.where(Document.case_id == case_id)
    if client_id is not None:
        query = query.where(Document.client_id == client_id)
    rows = await db.scalars(query.options(selectinload(Document.case), selectinload(Document.client), selectinload(Document.uploader)).order_by(Document.created_at.desc()))
    return [to_response(d) for d in rows.all()]


@router.get("/query", response_model=DocumentListResponse)
async def query_documents(
    document_id: int | None = None,
    search: str | None = None,
    case_id: int | None = None,
    client_id: int | None = None,
    category: str | None = None,
    file_type: str | None = None,
    uploaded_by: int | None = None,
    created_from: date | None = None,
    created_to: date | None = None,
    visibility: str | None = None,
    sort_by: str = "updated",
    page: int = 1,
    per_page: int = 10,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(role_guard(ALLOWED_STAFF)),
):
    if page < 1 or per_page < 1 or per_page > 100:
        raise HTTPException(status_code=422, detail="Invalid pagination")
    if visibility and visibility not in VALID_VISIBILITY:
        raise HTTPException(status_code=400, detail="Invalid document status")
    filters = [
        Document.organization_id == current_user.organization_id,
        or_(Document.case_id.is_(None), accessible_case_condition(current_user)),
    ]
    if document_id is not None:
        filters.append(Document.id == document_id)
    if case_id is not None:
        filters.append(Document.case_id == case_id)
    if client_id is not None:
        filters.append(or_(Document.client_id == client_id, Case.client_id == client_id))
    if category:
        filters.append(Document.category == category.strip())
    if file_type:
        filters.append(Document.file_type.ilike(f"%{file_type.strip()}%"))
    if uploaded_by is not None:
        filters.append(Document.uploaded_by == uploaded_by)
    if visibility:
        filters.append(Document.visibility == visibility)
    if created_from:
        filters.append(Document.created_at >= datetime.combine(created_from, datetime.min.time(), tzinfo=timezone.utc))
    if created_to:
        filters.append(Document.created_at < datetime.combine(created_to, datetime.min.time(), tzinfo=timezone.utc) + timedelta(days=1))
    if search and search.strip():
        term = f"%{search.strip()[:100]}%"
        filters.append(or_(Document.title.ilike(term), Document.file_name.ilike(term), Case.title.ilike(term), Client.name.ilike(term), cast(Document.id, String).ilike(term)))

    joins = select(Document).outerjoin(Case, Case.id == Document.case_id).outerjoin(Client, Client.id == func.coalesce(Document.client_id, Case.client_id))
    count_query = select(func.count(Document.id)).select_from(Document).outerjoin(Case, Case.id == Document.case_id).outerjoin(Client, Client.id == func.coalesce(Document.client_id, Case.client_id))
    total = int((await db.scalar(count_query.where(and_(*filters)))) or 0)
    order_map = {
        "name": Document.title.asc(),
        "created": Document.created_at.desc(),
        "case": Case.title.asc().nullslast(),
        "size": Document.file_size.desc().nullslast(),
        "updated": Document.updated_at.desc(),
    }
    rows = await db.scalars(
        joins.where(and_(*filters))
        .options(selectinload(Document.case), selectinload(Document.client), selectinload(Document.uploader))
        .order_by(order_map.get(sort_by, Document.updated_at.desc()), Document.id.desc())
        .offset((page - 1) * per_page).limit(per_page)
    )
    return DocumentListResponse(
        items=[to_response(document) for document in rows.all()], total=total, page=page,
        per_page=per_page, total_pages=max(1, (total + per_page - 1) // per_page),
    )


@router.get("/{document_id}", response_model=DocumentResponse)
async def get_document(
    document_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(role_guard(ALLOWED_STAFF)),
):
    doc = await get_org_document(db, document_id, current_user)
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    return to_response(doc)


@router.post("/{document_id}/onlyoffice/session", response_model=OnlyOfficeSessionResponse)
async def create_onlyoffice_session(
    document_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(role_guard(ALLOWED_STAFF)),
):
    document_server_url = (settings.onlyoffice_document_server_url or "").strip()
    if not document_server_url:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Online editor is not configured. Use basic editor or Replace File.",
        )

    doc = await get_org_document(db, document_id, current_user)
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    if not is_docx_document(doc):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="ONLYOFFICE editing currently supports DOCX only")

    backend_base_url = get_public_backend_base_url(request)
    file_token = build_internal_document_token(
        document_id=doc.id,
        organization_id=doc.organization_id,
        version=doc.version,
        purpose="onlyoffice_file",
    )
    callback_token = build_internal_document_token(
        document_id=doc.id,
        organization_id=doc.organization_id,
        version=doc.version,
        purpose="onlyoffice_callback",
    )
    file_url = f"{backend_base_url}/api/v1/documents/{doc.id}/onlyoffice/file?token={file_token}"
    callback_url = f"{backend_base_url}/api/v1/documents/{doc.id}/onlyoffice/callback?token={callback_token}"

    config = {
        "documentType": "word",
        "document": {
            "fileType": "docx",
            "title": doc.file_name,
            "url": file_url,
            "key": build_onlyoffice_document_key(doc),
            "permissions": {
                "edit": True,
                "download": True,
            },
        },
        "editorConfig": {
            "callbackUrl": callback_url,
            "mode": "edit",
            "user": {
                "id": str(current_user.id),
                "name": current_user.name or current_user.email or f"User {current_user.id}",
            },
            "customization": {
                "autosave": True,
            },
        },
    }

    return OnlyOfficeSessionResponse(
        document_id=doc.id,
        version=doc.version,
        document_server_url=normalize_base_url(document_server_url),
        editor_config=sign_onlyoffice_config(config),
        warning=DOCX_WARNING,
        notes=[
            "Edits saved from ONLYOFFICE create a new document version.",
            "The existing basic editor remains available as a fallback for DOCX text edits.",
        ],
    )


@router.get("/{document_id}/onlyoffice/file")
async def download_onlyoffice_document_file(
    document_id: int,
    token: str,
    db: AsyncSession = Depends(get_db),
):
    payload = decode_internal_document_token(token, expected_purpose="onlyoffice_file")
    doc = await get_org_document(db, document_id, int(payload["organization_id"]))
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    if int(payload["document_id"]) != doc.id or int(payload["version"]) != doc.version:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="ONLYOFFICE file token no longer matches the current version")
    if not is_docx_document(doc):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="ONLYOFFICE editing currently supports DOCX only")

    path = Path(doc.file_path)
    if not path.exists():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Stored file not found")
    return FileResponse(path=str(path), filename=doc.file_name, media_type=doc.file_type or DOCX_MIME_TYPE)


@router.get("/{document_id}/download")
async def download_document(
    document_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(role_guard(ALLOWED_STAFF)),
):
    doc = await get_org_document(db, document_id, current_user)
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    path = resolve_stored_file(doc.file_path, STORAGE_ROOT)
    return FileResponse(path=str(path), filename=doc.file_name, media_type=doc.file_type or "application/octet-stream")


@router.get("/{document_id}/view")
async def view_document(
    document_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(role_guard(ALLOWED_STAFF)),
):
    doc = await get_org_document(db, document_id, current_user)
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    path = resolve_stored_file(doc.file_path, STORAGE_ROOT)
    return FileResponse(
        path=str(path),
        filename=doc.file_name,
        media_type=doc.file_type or "application/octet-stream",
        content_disposition_type="inline",
    )


@router.patch("/{document_id}", response_model=DocumentResponse)
async def update_document(
    document_id: int,
    payload: DocumentUpdate,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(role_guard(ALLOWED_STAFF)),
):
    doc = await get_org_document(db, document_id, current_user)
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    updates = payload.model_dump(exclude_unset=True)
    if "visibility" in updates and updates["visibility"] not in VALID_VISIBILITY:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid visibility")
    for key, value in updates.items():
        setattr(doc, key, value)
    doc.updated_at = datetime.now(timezone.utc)

    if doc.case_id:
        await create_case_timeline_event(
            db,
            organization_id=current_user.organization_id,
            case_id=doc.case_id,
            actor_id=current_user.id,
            event_type="document_updated",
            title=f"Document updated: {doc.title}",
            metadata_json={"document_id": doc.id},
        )
    if updates.get("visibility") == "client_visible" and doc.case_id is not None:
        case = await db.scalar(select(Case).where(Case.id == doc.case_id, Case.organization_id == current_user.organization_id))
        if case:
            client = await db.scalar(select(Client).where(Client.id == case.client_id, Client.organization_id == current_user.organization_id))
            if client and client.user_id:
                await create_notification(
                    db,
                    organization_id=current_user.organization_id,
                    user_id=client.user_id,
                    type="document_shared",
                    title=f"Document shared: {doc.title}",
                    body="A document has been shared with you in the client portal.",
                    metadata_json={"document_id": doc.id, "case_id": case.id},
                )
                if client.email:
                    subject, html_body, text_body = build_document_shared_email(
                        client_name=client.name or "Client",
                        document_title=doc.title,
                        document_id=doc.id,
                    )
                    enqueue_email(background_tasks, to_email=client.email, subject=subject, html_body=html_body, text_body=text_body)

    await db.commit()
    await db.refresh(doc)
    return to_response(doc)


@router.post("/{document_id}/replace", response_model=DocumentResponse)
async def replace_document(
    document_id: int,
    file: UploadFile = File(...),
    notes: str | None = Form(default=None),
    request: Request = None,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(role_guard(ALLOWED_STAFF)),
):
    doc = await get_org_document(db, document_id, current_user)
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    previous_version = doc.version

    original_name = safe_original_name(file.filename or "")
    data = await file.read()
    if not data:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Empty file")
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="File exceeds upload size limit")

    file_path, _stored_name = persist_file(STORAGE_ROOT, current_user.organization_id, original_name, data)
    now = current_utc()

    version_row = archive_current_document_version(doc, current_user.id, now)
    db.add(version_row)

    doc.file_name = original_name
    doc.file_path = file_path
    doc.file_type = file.content_type
    doc.file_size = len(data)
    doc.version = previous_version + 1
    doc.uploaded_by = current_user.id
    doc.version_source = "replace"
    doc.version_note = notes.strip() if notes and notes.strip() else None
    doc.updated_at = now

    await log_audit_event(
        db,
        organization_id=current_user.organization_id,
        user_id=current_user.id,
        action="document_replaced",
        entity_type="document",
        entity_id=str(doc.id),
        description=f"Document replaced: {doc.title}",
        metadata_json={"previous_version": previous_version, "new_version": doc.version},
        ip_address=request.client.host if request and request.client else None,
        user_agent=request.headers.get("user-agent") if request else None,
    )
    if doc.case_id:
        await create_case_timeline_event(
            db,
            organization_id=current_user.organization_id,
            case_id=doc.case_id,
            actor_id=current_user.id,
            event_type="document_replaced",
            title=f"Document replaced: {doc.title}",
            metadata_json={"document_id": doc.id, "version": doc.version},
        )

    await db.commit()
    await db.refresh(doc)
    return to_response(doc)


@router.get("/{document_id}/editable-content", response_model=DocumentEditableContentResponse)
async def get_document_editable_content(
    document_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(role_guard(ALLOWED_STAFF)),
):
    doc = await get_org_document(db, document_id, current_user)
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    if not is_docx_document(doc):
        return DocumentEditableContentResponse(
            document_id=doc.id,
            file_type=doc.file_type,
            editable=False,
            mode=None,
            content="",
            warning=DOCX_WARNING,
            reason="DOCX editing only is supported right now. PDF editing will be added later.",
        )

    return DocumentEditableContentResponse(
        document_id=doc.id,
        file_type=doc.file_type or DOCX_MIME_TYPE,
        editable=True,
        mode="docx_text",
        content=extract_docx_text(doc.file_path),
        warning=DOCX_WARNING,
        reason=None,
    )


@router.post("/{document_id}/editable-content", response_model=DocumentResponse)
async def save_document_editable_content(
    document_id: int,
    payload: DocumentEditableContentUpdate,
    request: Request,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(role_guard(ALLOWED_STAFF)),
):
    doc = await get_org_document(db, document_id, current_user)
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    if not is_docx_document(doc):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Only DOCX documents can be edited in this workflow")

    previous_version = doc.version
    version_note = payload.version_note.strip() if payload.version_note and payload.version_note.strip() else None

    data = render_docx_bytes(payload.content)
    file_name = build_docx_version_name(doc)
    file_path, _stored_name = persist_file(STORAGE_ROOT, current_user.organization_id, file_name, data)
    now = current_utc()

    db.add(archive_current_document_version(doc, current_user.id, now))

    doc.file_name = file_name
    doc.file_path = file_path
    doc.file_type = DOCX_MIME_TYPE
    doc.file_size = len(data)
    doc.version = previous_version + 1
    doc.uploaded_by = current_user.id
    doc.version_source = "content_edit"
    doc.version_note = version_note
    doc.updated_at = now

    await log_audit_event(
        db,
        organization_id=current_user.organization_id,
        user_id=current_user.id,
        action="document_content_edited",
        entity_type="document",
        entity_id=str(doc.id),
        description=f"Document content edited: {doc.title}",
        metadata_json={"previous_version": previous_version, "new_version": doc.version, "source": "content_edit"},
        ip_address=request.client.host if request.client else None,
        user_agent=request.headers.get("user-agent"),
    )
    if doc.case_id:
        await create_case_timeline_event(
            db,
            organization_id=current_user.organization_id,
            case_id=doc.case_id,
            actor_id=current_user.id,
            event_type="document_content_edited",
            title=f"Document edited: {doc.title}",
            metadata_json={"document_id": doc.id, "version": doc.version},
        )

    await db.commit()
    await db.refresh(doc)
    return to_response(doc)


@router.post("/{document_id}/onlyoffice/callback")
async def handle_onlyoffice_callback(
    document_id: int,
    request: Request,
    token: str,
    db: AsyncSession = Depends(get_db),
):
    payload = decode_internal_document_token(token, expected_purpose="onlyoffice_callback")
    doc = await get_org_document(db, document_id, int(payload["organization_id"]))
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    if int(payload["document_id"]) != doc.id:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="ONLYOFFICE callback token does not match this document")
    if not is_docx_document(doc):
        return ONLYOFFICE_CALLBACK_SUCCESS

    body = await request.json()
    callback_payload = get_onlyoffice_callback_payload(request, body)
    status_code = int(callback_payload.get("status") or 0)
    if status_code not in ONLYOFFICE_SAVE_STATUSES:
        return ONLYOFFICE_CALLBACK_SUCCESS

    expected_key = build_onlyoffice_document_key(doc)
    if callback_payload.get("key") != expected_key:
        return ONLYOFFICE_CALLBACK_SUCCESS
    if int(payload["version"]) != doc.version:
        return ONLYOFFICE_CALLBACK_SUCCESS

    file_url = callback_payload.get("url")
    if not file_url or not isinstance(file_url, str):
        return ONLYOFFICE_CALLBACK_SUCCESS

    async with httpx.AsyncClient(timeout=60.0) as client:
        response = await client.get(file_url)
        response.raise_for_status()
        data = response.content

    if not data:
        return ONLYOFFICE_CALLBACK_SUCCESS
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Edited file exceeds upload size limit")

    current_file_bytes = Path(doc.file_path).read_bytes()
    if current_file_bytes == data:
        return ONLYOFFICE_CALLBACK_SUCCESS

    actor_user_id = get_onlyoffice_actor_user_id(callback_payload, doc.uploaded_by)
    now = current_utc()
    db.add(archive_current_document_version(doc, actor_user_id, now))

    file_name = build_docx_version_name(doc)
    file_path, _stored_name = persist_file(STORAGE_ROOT, doc.organization_id, file_name, data)
    previous_version = doc.version
    doc.file_name = file_name
    doc.file_path = file_path
    doc.file_type = DOCX_MIME_TYPE
    doc.file_size = len(data)
    doc.version = previous_version + 1
    doc.uploaded_by = actor_user_id
    doc.version_source = "onlyoffice_edit"
    doc.version_note = build_onlyoffice_callback_note(callback_payload)
    doc.updated_at = now

    await log_audit_event(
        db,
        organization_id=doc.organization_id,
        user_id=actor_user_id,
        action="document_onlyoffice_edited",
        entity_type="document",
        entity_id=str(doc.id),
        description=f"Document edited in ONLYOFFICE: {doc.title}",
        metadata_json={"previous_version": previous_version, "new_version": doc.version, "source": "onlyoffice_edit"},
        ip_address=request.client.host if request.client else None,
        user_agent=request.headers.get("user-agent"),
    )
    if doc.case_id:
        await create_case_timeline_event(
            db,
            organization_id=doc.organization_id,
            case_id=doc.case_id,
            actor_id=actor_user_id,
            event_type="document_onlyoffice_edited",
            title=f"Document edited: {doc.title}",
            metadata_json={"document_id": doc.id, "version": doc.version, "source": "onlyoffice_edit"},
        )

    await db.commit()
    return ONLYOFFICE_CALLBACK_SUCCESS


@router.get("/{document_id}/versions", response_model=list[DocumentVersionResponse])
async def list_document_versions(
    document_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(role_guard(ALLOWED_STAFF)),
):
    doc = await get_org_document(db, document_id, current_user)
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    rows = await db.scalars(
        select(DocumentVersion)
        .where(DocumentVersion.document_id == document_id, DocumentVersion.organization_id == current_user.organization_id)
        .order_by(DocumentVersion.version_number.desc())
    )
    return [to_version_response(row) for row in rows.all()]


@router.get("/{document_id}/versions/{version_id}/download")
async def download_document_version(
    document_id: int,
    version_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(role_guard(ALLOWED_STAFF)),
):
    doc = await get_org_document(db, document_id, current_user)
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    version = await db.scalar(
        select(DocumentVersion).where(
            DocumentVersion.id == version_id,
            DocumentVersion.document_id == document_id,
            DocumentVersion.organization_id == current_user.organization_id,
        )
    )
    if not version:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Version not found")
    path = Path(version.file_path)
    if not path.exists():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Stored file not found")
    return FileResponse(path=str(path), filename=version.file_name, media_type=version.file_type or "application/octet-stream")


@router.delete("/{document_id}")
async def delete_document(
    document_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(role_guard(ALLOWED_STAFF)),
):
    doc = await get_org_document(db, document_id, current_user)
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    if doc.case_id:
        await create_case_timeline_event(
            db,
            organization_id=current_user.organization_id,
            case_id=doc.case_id,
            actor_id=current_user.id,
            event_type="document_deleted",
            title=f"Document deleted: {doc.title}",
            metadata_json={"document_id": doc.id, "file_name": doc.file_name},
        )
    await log_audit_event(
        db,
        organization_id=current_user.organization_id,
        user_id=current_user.id,
        action="document_deleted",
        entity_type="document",
        entity_id=str(doc.id),
        description=f"Document deleted: {doc.title}",
        metadata_json={"case_id": doc.case_id},
        ip_address=request.client.host if request.client else None,
        user_agent=request.headers.get("user-agent"),
    )

    path = Path(doc.file_path)
    if path.exists():
        path.unlink(missing_ok=True)

    await db.delete(doc)
    await db.commit()
    return {"ok": True}
